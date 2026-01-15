"DOCX document interface."

import datetime

import docx
import docx.oxml
import docx.shared
import docx.styles.style

from .base_document import *

__all__ = ["Document"]

MAX_LEVEL = 6
HEADER_SPACE_BEFORE = 14
HEADER_SPACE_AFTER = 6
NORMAL_FONT = "Helvetica"
NORMAL_FONT_SIZE = 12
NORMAL_LINE_SPACING = 18
QUOTE_FONT = "Times New Roman"
QUOTE_FONT_SIZE = 14
QUOTE_INDENT = 24
# CODE_FONT = "Courier"
# CODE_FONT_SIZE = 11
# CODE_LINE_SPACING = 12
# CODE_INDENT = 10
TITLE_PAGE_SPACER = 80
HYPERLINK_COLOR = (0x05, 0x63, 0xC1)
EMDASH = "\u2014"


class Document(BaseDocument):
    "DOCX document interface."

    def __init__(self, **kwargs):
        super().__init__(**kwargs)

        self.indexed = {}
        self.toc_table = None

        self.docx = docx.Document()

        # Set the default document-wide language.
        # From https://stackoverflow.com/questions/36967416/how-can-i-set-the-language-in-text-with-python-docx
        if self.language:
            styles_element = self.docx.styles.element
            rpr_default = styles_element.xpath("./w:docDefaults/w:rPrDefault/w:rPr")[0]
            lang_default = rpr_default.xpath("w:lang")[0]
            lang_default.set(docx.oxml.shared.qn("w:val"), self.language)

        # Set to A4 page size. The section here is an instance of the docx Section class.
        section = self.docx.sections[-1]
        section.page_height = docx.shared.Mm(297)
        section.page_width = docx.shared.Mm(210)
        section.left_margin = docx.shared.Mm(25.4)  # 1 inch
        section.right_margin = docx.shared.Mm(25.4)
        section.top_margin = docx.shared.Mm(25.4)
        section.bottom_margin = docx.shared.Mm(25.4)
        section.header_distance = docx.shared.Mm(12.7)  # 0.5 inch
        section.footer_distance = docx.shared.Mm(12.7)

        # Modify styles.
        style = self.docx.styles["Title"]
        style.font.color.rgb = docx.shared.RGBColor(0, 0, 0)

        for level in range(1, MAX_LEVEL + 1):
            style = self.docx.styles[f"Heading {level}"]
            style.paragraph_format.space_before = docx.shared.Pt(HEADER_SPACE_BEFORE)
            style.paragraph_format.space_after = docx.shared.Pt(HEADER_SPACE_AFTER)
            style.font.color.rgb = docx.shared.RGBColor(0, 0, 0)

        style = self.docx.styles["Normal"]
        style.font.name = NORMAL_FONT
        style.font.size = docx.shared.Pt(NORMAL_FONT_SIZE)
        style.paragraph_format.line_spacing = docx.shared.Pt(NORMAL_LINE_SPACING)

        style = self.docx.styles["Quote"]  # Quote blocks.
        style.font.name = QUOTE_FONT
        style.font.size = docx.shared.Pt(QUOTE_FONT_SIZE)
        style.font.italic = False
        style.paragraph_format.left_indent = docx.shared.Pt(QUOTE_INDENT)
        style.paragraph_format.right_indent = docx.shared.Pt(QUOTE_INDENT)

        # style = self.docx.styles["macro"]  # Code blocks.
        # style.font.name = CODE_FONT
        # style.font.size = docx.shared.Pt(CODE_FONT_SIZE)
        # style.paragraph_format.line_spacing = docx.shared.Pt(CODE_LINE_SPACING)
        # style.paragraph_format.left_indent = docx.shared.Pt(CODE_INDENT)

        style = self.docx.styles["Body Text 2"]  # TOC entry.
        style.font.name = NORMAL_FONT
        style.font.size = docx.shared.Pt(NORMAL_FONT_SIZE)
        style.paragraph_format.space_before = docx.shared.Mm(1)
        style.paragraph_format.space_after = docx.shared.Mm(1)
        style.paragraph_format.line_spacing = docx.shared.Mm(1)

        style = self.docx.styles["Body Text 3"]  # TOC page.
        style.font.name = NORMAL_FONT
        style.font.size = docx.shared.Pt(NORMAL_FONT_SIZE)
        style.paragraph_format.space_before = docx.shared.Mm(1)
        style.paragraph_format.space_after = docx.shared.Mm(1)
        style.paragraph_format.line_spacing = docx.shared.Mm(1)
        style.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.RIGHT

        # Add new styles.
        style = self.docx.styles.add_style(  # Hyperlink.
            "Hyperlink", docx.enum.style.WD_STYLE_TYPE.CHARACTER, True
        )
        style.base_style = self.docx.styles["Normal"]
        style.unhide_when_used = True
        style.font.color.rgb = docx.shared.RGBColor(*HYPERLINK_COLOR)
        style.font.underline = True

        style = self.docx.styles.add_style(
            "List Number Quote", docx.enum.style.WD_STYLE_TYPE.CHARACTER, True
        )
        style.base_style = self.docx.styles["List Number"]
        style.font.name = QUOTE_FONT
        style.font.size = docx.shared.Pt(QUOTE_FONT_SIZE)
        style.unhide_when_used = True

        style = self.docx.styles.add_style(
            "List Bullet Quote", docx.enum.style.WD_STYLE_TYPE.CHARACTER, True
        )
        style.base_style = self.docx.styles["List Bullet"]
        style.font.name = QUOTE_FONT
        style.font.size = docx.shared.Pt(QUOTE_FONT_SIZE)
        style.unhide_when_used = True

        style = self.docx.styles.add_style(
            "List Continue Quote", docx.enum.style.WD_STYLE_TYPE.CHARACTER, True
        )
        style.base_style = self.docx.styles["List Continue"]
        style.font.name = QUOTE_FONT
        style.font.size = docx.shared.Pt(QUOTE_FONT_SIZE)
        style.unhide_when_used = True

        # Set Dublin core metadata.
        if self.authors:
            self.docx.core_properties.author = ", ".join(self.authors)
        self.docx.core_properties.created = datetime.datetime.now(tz=datetime.UTC)
        if self.language:
            self.docx.core_properties.language = self.language

        # Display page number in the header.
        # https://stackoverflow.com/questions/56658872/add-page-number-using-python-docx
        paragraph = self.docx.sections[-1].header.paragraphs[0]
        paragraph.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.RIGHT
        run = paragraph.add_run()
        fldChar1 = docx.oxml.OxmlElement("w:fldChar")
        fldChar1.set(docx.oxml.ns.qn("w:fldCharType"), "begin")
        instrText = docx.oxml.OxmlElement("w:instrText")
        instrText.set(docx.oxml.ns.qn("xml:space"), "preserve")
        instrText.text = "PAGE"
        fldChar2 = docx.oxml.OxmlElement("w:fldChar")
        fldChar2.set(docx.oxml.ns.qn("w:fldCharType"), "end")
        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)

        if self.title:
            self.docx.add_heading(self.title, 0)
        if self.subtitle:
            self.docx.add_heading(self.subtitle, 1)
        if self.authors:
            self.docx.add_heading(", ".join(self.authors), 2)
        if self.version:
            self.docx.add_paragraph(self.version)
        p = self.docx.add_paragraph(EMDASH * 35, style="Normal")
        p.paragraph_format.space_after = docx.shared.Pt(TITLE_PAGE_SPACER)

    def initialize_toc(self):
        if not self.toc_level:
            return
        if self.toc_table:
            return
        self.docx.add_page_break()
        self.docx.add_heading(self.toc_title, 1)
        self.toc_table = self.docx.add_table(
            rows=0, cols=0, style=self.docx.styles["Normal Table"]
        )
        self.toc_table.add_column(docx.shared.Mm(140))
        self.toc_table.add_column(docx.shared.Mm(20))

    def paragraph(self, text=None, thematic_break=False):
        """Create a new paragraph, add the text (if any) to it and return it.
        Optionally add a thematic break before it.
        """
        if thematic_break:
            self.thematic_break()
        paragraph = Paragraph(self)
        if text:
            paragraph.add(text)
        return paragraph

    def thematic_break(self):
        p = self.docx.add_paragraph(EMDASH * 20, style="Normal")
        p.paragraph_format.space_before = docx.shared.Mm(8)
        p.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER

    def quote(self, text=None):
        "Create a new quotation paragraph, add the text (if any) to it and return it."
        paragraph = Quote(self)
        if text:
            paragraph.add(text)
        return paragraph

    def section(self, title, subtitle=None):
        "Create a new subsection, which is a context that increments the section level."
        return Section(self, title, subtitle=subtitle)

    def pagebreak(self):
        self.docx.add_page_break()

    def ordered_list(self):
        return List(self, ordered=True)

    def unordered_list(self):
        return List(self, ordered=False)

    def write(self, filepath):
        self.docx.save(filepath)


class Section(BaseSection):

    def __init__(self, document, title, subtitle=None):
        super().__init__(document, title, subtitle=subtitle)
        self.document.initialize_toc()

    def __enter__(self):
        super().__enter__()
        self.document.docx.add_heading(self.title, level=self.level)
        if self.subtitle:
            self.document.docx.add_heading(self.subtitle, level=self.level + 1)
        if self.level <= self.document.toc_level:
            cells = self.document.toc_table.add_row().cells
            cells[0].add_paragraph(self.title, style="Body Text 2")
            cells[0].paragraphs[0].paragraph_format.left_indent = docx.shared.Mm(
                3 * (self.level - 1)
            )
            self.page_number_paragraph = cells[1].add_paragraph(
                str(self.document.page["docx"]), style="Body Text 3"
            )
        return self

    def set_page(self, **kwargs):
        super().set_page(**kwargs)
        try:
            paragraph = self.page_number_paragraph
        except AttributeError:
            pass
        else:
            paragraph.text = str(self.document.page["docx"])

    def output_footnotes(self, title="Footnotes"):
        "Output the footnotes to the section."
        if not self.document.footnotes:
            return
        with self.document.no_numbers():
            self.document.thematic_break()
            self.document.docx.add_heading(title, level=self.level + 1)
            self.document.output_footnotes_list()


class Paragraph(BaseParagraph):

    STYLESHEETNAME = "Normal"

    def __init__(self, document):
        super().__init__(document)
        self.paragraph = document.docx.add_paragraph(style=self.STYLESHEETNAME)
        self.line_started = False
        self._bold = 0
        self._italic = 0
        self._underline = 0
        self._subscript = 0
        self._superscript = 0
        if document.paragraph_numbers:
            self.add(f"({document.paragraphs_count})")

    def add(self, text, raw=False):
        """Add the text to the paragraph.
        - Exchange newlines for blanks.
        - Remove superfluous blanks.
        - Prepend a blank, if 'raw' is False.
        - Return the paragraph.
        """
        assert isinstance(text, str)
        lines = list(text.split())
        if raw or not self.line_started:
            result = []
        else:
            result = [" "]
        for line in lines[:-1]:
            result.append(line)
            result.append(" ")
        try:
            result.append(lines[-1])
        except IndexError:
            pass
        self.set_font_style(self.paragraph.add_run("".join(result)))
        self.line_started = True
        return self

    def linebreak(self):
        "Add a line break. Return the paragraph."
        self.paragraph.add_run("\n")
        self.line_started = False
        return self

    def emdash(self, raw=False):
        "Add an emdash character. Return the paragraph."
        if not raw and self.line_started:
            self.set_font_style(self.paragraph.add_run(" "))
        self.set_font_style(self.paragraph.add_run(EMDASH))
        self.line_started = True
        return self

    def indexed(self, text, canonical=None, raw=False):
        "Add an indexed term, optionally with its canonical term. Return the paragraph."
        if not raw and self.line_started:
            self.set_font_style(self.paragraph.add_run(" "))
        with self.underline():
            self.raw(text)
        self.line_started = True
        self.document.add_indexed(canonical or text)
        return self

    def link(self, href, text=None, raw=False):
        "Add a hyperlink, optionally with a text to display. Return the paragraph."
        if not raw and self.line_started:
            self.set_font_style(self.paragraph.add_run(" "))

        # https://github.com/python-openxml/python-docx/issues/74#issuecomment-261169410
        # This works in 'writethatbook', but not here??

        # This gets access to the document.xml.rels file and gets a new relation id value
        part = self.paragraph.part
        r_id = part.relate_to(
            href, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True
        )

        # Create the w:hyperlink tag and add needed values
        hyperlink = docx.oxml.shared.OxmlElement("w:hyperlink")
        hyperlink.set(docx.oxml.shared.qn("r:id"), r_id)

        # Create a w:r element and a new w:rPr element and join together
        new_run = docx.oxml.shared.OxmlElement("w:r")
        rPr = docx.oxml.shared.OxmlElement("w:rPr")
        new_run.append(rPr)

        # Style and add text to the w:r element
        new_run.text = text or href
        new_run.style = "Hyperlink"
        hyperlink.append(new_run)

        self.paragraph._p.append(hyperlink)
        self.line_started = True
        return self

    def comment(self, text):
        try:
            author = self.document.authors[0]
        except IndexError:
            author = None
        self.document.docx.add_comment(
            runs=self.paragraph.runs, text=text, author=author
        )

    @contextmanager
    def bold(self):
        try:
            self._bold += 1
            yield self
        finally:
            self._bold -= 1

    @contextmanager
    def italic(self):
        try:
            self._italic += 1
            yield self
        finally:
            self._italic -= 1

    @contextmanager
    def underline(self):
        try:
            self._underline += 1
            yield self
        finally:
            self._underline -= 1

    @contextmanager
    def subscript(self):
        try:
            self._subscript += 1
            yield self
        finally:
            self._subscript -= 1

    @contextmanager
    def superscript(self):
        try:
            self._superscript += 1
            yield self
        finally:
            self._superscript -= 1

    def set_font_style(self, run):
        if self._bold:
            run.font.bold = True
        if self._italic:
            run.font.italic = True
        if self._underline:
            run.font.underline = True
        if self._subscript:
            run.font.subscript = True
        if self._superscript:
            run.font.superscript = True


class Quote(Paragraph):

    STYLESHEETNAME = "Quote"


class List(BaseList):
    "List class; a context manager."

    def __init__(self, document, ordered=False, level=1):
        super().__init__(document, ordered=ordered)
        self.level = level

    def item(self):
        return ListItem(self)


class ListItem(BaseListItem):
    "List item class; a context manager."

    def __enter__(self):
        self.first_paragraph = True
        return self

    def __exit__(self, *exc):
        pass

    def paragraph(self, text=None):
        if self.first_paragraph:
            if self.list.ordered:
                paragraph = OrderedListItemParagraph(
                    self.list.document, self.list.level
                )
            else:
                paragraph = UnorderedListItemParagraph(
                    self.list.document, self.list.level
                )
            self.first_paragraph = False
        else:
            paragraph = ContinueListItemParagraph(self.list.document, self.list.level)
        if text:
            paragraph.add(text)
        return paragraph

    def quote(self, text=None):
        if self.first_paragraph:
            if self.list.ordered:
                quote = OrderedListItemQuote(self.list.document, self.list.level)
            else:
                quote = UnorderedListItemQuote(self.list.document, self.list.level)
            self.first_quote = False
        else:
            quote = ContinueListItemQuote(self.list.document, self.list.level)
        if text:
            quote.add(text)
        return quote

    def ordered_list(self):
        return List(self.list.document, ordered=True, level=self.list.level + 1)

    def unordered_list(self):
        return List(self.list.document, ordered=False, level=self.list.level + 1)


class ListItemParagraph(Paragraph):

    def __init__(self, document, level):
        if level > 1:
            # Note: class level attribute is overridden by an instance level attribute.
            self.STYLESHEETNAME = f"{self.STYLESHEETNAME} {level}"
        super().__init__(document)


class OrderedListItemParagraph(ListItemParagraph):

    STYLESHEETNAME = "List Number"


class UnorderedListItemParagraph(ListItemParagraph):

    STYLESHEETNAME = "List Bullet"


class ContinueListItemParagraph(ListItemParagraph):

    STYLESHEETNAME = "List Continue"


class OrderedListItemQuote(ListItemParagraph):

    STYLESHEETNAME = "List Number Quote"


class UnorderedListItemQuote(ListItemParagraph):

    STYLESHEETNAME = "List Bullet Quote"


class ContinueListItemQuote(ListItemParagraph):

    STYLESHEETNAME = "List Continue Quote"
